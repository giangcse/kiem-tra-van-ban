import os
import docx
from fastapi import FastAPI, File, UploadFile
from pydantic import BaseModel
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from typing import List

app = FastAPI()

class FileUploadRequest(BaseModel):
    files: List[UploadFile] = File(...)

def compare_docx_files(doc1, doc2):
    for table1, table2 in zip(doc1.tables, doc2.tables):
        if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(table2.columns):
            print("Tables have different sizes")
            continue
        for i, row1 in enumerate(table1.rows):
            row2 = table2.rows[i]
            for j, cell1 in enumerate(row1.cells):
                cell2 = row2.cells[j]
                if cell1.text != cell2.text:
                    for paragraph in cell1.paragraphs:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    for paragraph in cell2.paragraphs:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    cell1.text = f"{cell1.text} ({cell2.text})"
                    cell1.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    return doc1

@app.post("/compare-docx-files/")
async def compare_docx_files_route(req: FileUploadRequest):
    file1 = req.files[0]
    file2 = req.files[1]
    file1_path = os.path.join(os.getcwd(), file1.filename)
    file2_path = os.path.join(os.getcwd(), file2.filename)
    with open(file1_path, "wb") as f1, open(file2_path, "wb") as f2:
        f1.write(await file1.read())
        f2.write(await file2.read())
    doc1 = docx.Document(file1_path)
    doc2 = docx.Document(file2_path)
    compare_docx_files(doc1, doc2)
    result_file_path = os.path.join(os.getcwd(), "result.docx")
    doc1.save(result_file_path)
    with open(result_file_path, "rb") as f:
        result_file_content = f.read()
    os.remove(file1_path)
    os.remove(file2_path)
    os.remove(result_file_path)
    return {"result": result_file_content}

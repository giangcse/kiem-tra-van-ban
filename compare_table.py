import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX

def compare_docx_files(doc1_filename, doc2_filename, output_filename):
    # Load the first document
    doc1 = docx.Document(doc1_filename)

    # Load the second document
    doc2 = docx.Document(doc2_filename)

    # Compare the tables in each document
    for table1, table2 in zip(doc1.tables, doc2.tables):
        # Compare the first header in each table
        if table1.rows[0].cells[0].text != table2.rows[0].cells[0].text:
            cell1 = table1.rows[0].cells[0]
            cell2 = table2.rows[0].cells[0]
            cell1.paragraphs[0].runs[0].font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            cell2.paragraphs[0].runs[0].font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

        # Compare the second header in each table
        if table1.rows[1].cells[0].text != table2.rows[1].cells[0].text:
            cell1 = table1.rows[1].cells[0]
            cell2 = table2.rows[1].cells[0]
            cell1.paragraphs[0].runs[0].font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            cell2.paragraphs[0].runs[0].font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

        # Compare the data rows in each table
        for row1, row2 in zip(table1.rows[2:], table2.rows[2:]):
            for cell1, cell2 in zip(row1.cells, row2.cells):
                for para1, para2 in zip(cell1.paragraphs, cell2.paragraphs):
                    for run1, run2 in zip(para1.runs, para2.runs):
                        for sent1, sent2 in zip(run1.text.split('. '), run2.text.split('. ')):
                            if sent1 != sent2:
                                run1.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                                run2.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
    
    
    for i in range(min(len(doc1.paragraphs), len(doc2.paragraphs))):
        para1 = doc1.paragraphs[i]
        para2 = doc2.paragraphs[i]

        if para1.text != para2.text:
            # mark the difference in para1
            for run in para1.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            # mark the difference in para2
            for run in para2.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # Save the modified document
    doc1.save(output_filename)

compare_docx_files('tai_lieu_1.docx', 'tai_lieu_2.docx', 'tai_lieu_so_sanh.docx')
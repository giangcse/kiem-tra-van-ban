import pandas as pd
from docx import Document
from underthesea import sent_tokenize, text_normalize

def read_docx_table(document, table_num=1, n_headers=1):
    table = document.tables[table_num-1]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    if n_headers == 1:
        df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
    elif n_headers==2:
        outside_col, inside_col = df.iloc[0], df.iloc[1]
        hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col, inside_col)))
        df = pd.DataFrame(data, columns=hier_index).drop(df.index[[0, 1]]).reset_index(drop=True)
    elif n_headers>2:
        print("Not supported")
        df = pd.DataFrame()
    return df

def docx_to_elastic(filename):
    document = Document(filename)
    try:
        paragraphs = document.paragraphs
        tables = document.tables
        bulkData = []
        for para in paragraphs:
            for sentence in sent_tokenize(text_normalize(para.text)):
                bulkData.append(
                    {
                        "filename": str(filename),
                        "sentence": sentence,
                        "type": "sentence"
                    }
                )
        for table in tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text)
                bulkData.append(
                    {
                        "filename": str(filename),
                        "sentence": str('-'.join(cells)),
                        "type": "table"
                    }
                )
        return bulkData
    except Exception as e:
        return e
    

from elastic_bulk import bulk_data
bulkData = docx_to_elastic('tai_lieu.docx')
print(bulk_data(bulkData, "test001"))


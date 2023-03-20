import difflib
import docx

# Đọc file đầu tiên
doc1 = docx.Document("doc1.docx")
text1 = "\n".join([para.text for para in doc1.paragraphs])

# Đọc file thứ hai
doc2 = docx.Document("doc2.docx")
text2 = "\n".join([para.text for para in doc2.paragraphs])

# Tạo bản sao của file đầu tiên
doc1_copy = docx.Document("doc1.docx")

# So sánh và tô sáng sự khác biệt
d = difflib.Differ()
diff = d.compare(text1.splitlines(), text2.splitlines())
for line in diff:
    if line.startswith("- "):
        # Tìm kiếm đoạn văn bản trong file đầu tiên
        for para in doc1_copy.paragraphs:
            if line[2:] in para.text:
                # Tô sáng phần khác biệt
                for run in para.runs:
                    if line[2:] in run.text:
                        run.font.color.rgb = docx.shared.RGBColor(0xFF, 0xFF, 0x00) # tô sáng phần khác biệt bằng màu vàng

# Lưu file mới
doc1_copy.save("file1_copy.docx")

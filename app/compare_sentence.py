import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX

def compare_docs(file1_path, file2_path, output_path):
    doc1 = docx.Document(file1_path)
    doc2 = docx.Document(file2_path)

    for i in range(min(len(doc1.paragraphs), len(doc2.paragraphs))):
        para1 = doc1.paragraphs[i]
        para2 = doc2.paragraphs[i]

        if para1.text != para2.text:
            # mark the difference in para1
            for run in para1.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            # mark the difference in para2
            for run in para2.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for i in range(min(len(doc1.tables), len(doc2.tables))):
        table1 = doc1.tables[i]
        table2 = doc2.tables[i]
        if table1._element.xml != table2._element.xml:
            # mark the difference in table1
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER

            # mark the difference in table2
            for row in table2.rows:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc1.save(output_path)

compare_docs('doc1.docx', 'doc2.docx', 'tai_lieu_so_sanh.docx')